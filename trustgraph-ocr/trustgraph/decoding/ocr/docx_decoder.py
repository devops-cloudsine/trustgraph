"""
DOCX decoder: Extracts text, tables, and images from DOCX documents using python-docx.
Images are sent to vLLM for description. All content is output as structured text segments
suitable for RAG processing.

Key features:
- Uses python-docx (free, open-source, no license required)
- Fallback to zipfile extraction when python-docx fails (handles corrupt files)
- Concurrent vLLM requests for faster image processing
- Image deduplication by MD5 hash to avoid redundant LLM calls
- Robust error handling that continues processing even if parts fail
- Extracts text, tables, headers/footers, and images
- Images stored in MinIO for persistence

Reference: https://python-docx.readthedocs.io/
"""

import asyncio
import base64
import concurrent.futures
import hashlib
import logging
import os
import re
import urllib.parse
import zipfile
from io import BytesIO
from pathlib import Path
import requests
import json
from typing import Dict, List, Any, Optional, Set, Tuple

from ... schema import Document as TGDocument, TextDocument
from ... schema import Triples, Triple, Value, Metadata
from ... schema import EntityContext, EntityContexts
from ... base import FlowProcessor, ConsumerSpec, ProducerSpec
from .minio_storage import MinioStorage, get_minio_storage

# ---> rdf.py constants for image context triples
IMAGE_CONTEXT = "http://trustgraph.ai/ns/image-context"
IMAGE_SOURCE = "http://trustgraph.ai/ns/image-source"
TRUSTGRAPH_ENTITIES = "http://trustgraph.ai/e/"

# Module logger
logger = logging.getLogger(__name__)

# Try to import python-docx
try:
    from docx import Document as DocxDocument
    from docx.opc.exceptions import PackageNotFoundError
    from docx.shared import Inches, Pt
    DOCX_AVAILABLE = True
except ImportError as e:
    DOCX_AVAILABLE = False
    DocxDocument = None
    logger.warning(f"python-docx not available: {e}")

# Try to import PIL for image handling
try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

# Enable DEBUG level for this module when VLLM_LOGGING_LEVEL=DEBUG is set
if os.environ.get("VLLM_LOGGING_LEVEL", "").upper() == "DEBUG":
    logger.setLevel(logging.DEBUG)

default_ident = "docx-decoder"

# Performance tuning constants
MAX_CONCURRENT_VLLM_REQUESTS = 4  # Max parallel vLLM requests
VLLM_TIMEOUT_SECONDS = 180  # Timeout for vLLM requests
MAX_IMAGE_SIZE_MB = 10  # Skip images larger than this


class Processor(FlowProcessor):

    def __init__(self, **params):
        id = params.get("id", default_ident)
        self.vllm_api_url = params.get(
            "vllm_api_url",
            "http://vllm:8000/v1/chat/completions"
        )
        self.vllm_model = params.get(
            "vllm_model",
            "Qwen/Qwen3-VL-4B-Instruct"
        )
        
        # Initialize MinIO storage for image persistence
        self.minio_storage = get_minio_storage()

        super(Processor, self).__init__(
            **params | {"id": id}
        )

        self.register_specification(
            ConsumerSpec(
                name="input",
                schema=TGDocument,
                handler=self.on_message,
            )
        )

        self.register_specification(
            ProducerSpec(
                name="output",
                schema=TextDocument,
            )
        )

        # ---> on_message > [triples output] > emit IMAGE_CONTEXT/IMAGE_SOURCE triples for Graph RAG
        self.register_specification(
            ProducerSpec(
                name="triples",
                schema=Triples,
            )
        )

        # ---> on_message > [entity-contexts output] > emit EntityContexts for graph embeddings (enables graph-retrieval)
        self.register_specification(
            ProducerSpec(
                name="entity-contexts",
                schema=EntityContexts,
            )
        )

        logger.info("DOCX decoder initialized (python-docx + zipfile fallback, concurrent vLLM)")

    # // ---> Pulsar consumer(input) > [on_message] > extract text/tables/images, vLLM describe -> flow('output')
    async def on_message(self, msg, consumer, flow):
        logger.info("DOCX message received for extraction")

        v = msg.value()
        doc_id = v.metadata.id
        logger.info(f"Processing {doc_id}...")

        blob = base64.b64decode(v.data)

        # Check if this is a DOCX file
        content_type = getattr(v, "content_type", None)
        if not self._is_docx(blob, content_type):
            logger.info(f"Skipping non-DOCX file: {doc_id} (content_type: {content_type})")
            return

        # Initialize MinIO storage
        if not self.minio_storage.initialize():
            logger.error("Failed to initialize MinIO storage - cannot process DOCX")
            return

        # Save original DOCX document to MinIO
        docx_filename = f"{self._safe_id(doc_id)}.docx"
        doc_object_name = self.minio_storage.save_document(
            doc_id=doc_id,
            filename=docx_filename,
            document_data=blob,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        if doc_object_name:
            logger.info(f"Saved original DOCX to MinIO: {doc_object_name}")
        else:
            logger.warning("Failed to save original DOCX to MinIO, continuing with processing")

        # Track processed image hashes for deduplication
        processed_hashes: Set[str] = set()
        all_images: List[Dict[str, Any]] = []
        structured_data = None
        docx_success = False

        # Try python-docx extraction (loads from BytesIO)
        if DOCX_AVAILABLE:
            docx_success, structured_data, all_images = self._try_docx_extraction(
                blob, doc_id, processed_hashes
            )

        # Fallback to zipfile extraction if python-docx failed
        if not docx_success:
            logger.info("Using zipfile fallback extraction")
            fallback_images = self._extract_images_from_zip(blob, doc_id, processed_hashes)
            for ix, img_info in enumerate(fallback_images):
                all_images.append({
                    "data": img_info["data"],
                    "content_type": img_info["content_type"],
                    "context": f"Image {ix + 1} ({img_info['original_name']})",
                    "paragraph_index": None
                })
            structured_data = {
                "document": {
                    "paragraphs": [],
                    "tables": [],
                    "properties": {},
                    "note": "Extracted via zipfile fallback"
                }
            }
            logger.info(f"Fallback: Found {len(fallback_images)} images")

        # Send document header first
        if structured_data:
            filename = getattr(v.metadata, 'id', 'document') + '.docx'
            header = self._format_document_header(structured_data, filename)
            if header.strip():
                logger.info(f"Sending document header to RAG ({len(header)} chars)")
                r = TextDocument(
                    metadata=v.metadata,
                    text=header.encode("utf-8"),
                )
                await flow("output").send(r)

        # Send text content in segments
        if structured_data and structured_data["document"]["paragraphs"]:
            text_content = self._format_document_text(structured_data)
            if text_content.strip():
                logger.info(f"Sending document text to RAG ({len(text_content)} chars)")
                r = TextDocument(
                    metadata=v.metadata,
                    text=text_content.encode("utf-8"),
                )
                await flow("output").send(r)

        # Send tables if any
        if structured_data and structured_data["document"]["tables"]:
            tables_text = self._format_tables(structured_data["document"]["tables"])
            if tables_text.strip():
                logger.info(f"Sending tables to RAG ({len(tables_text)} chars)")
                r = TextDocument(
                    metadata=v.metadata,
                    text=tables_text.encode("utf-8"),
                )
                await flow("output").send(r)

        # Process images with concurrent vLLM requests
        # Collect triples and entity contexts for Graph RAG
        image_triples = []
        entity_contexts = []  # For graph embeddings
        doc_uri = TRUSTGRAPH_ENTITIES + urllib.parse.quote(doc_id)
        
        if all_images:
            logger.info(f"Processing {len(all_images)} unique images with concurrent vLLM requests")
            
            # Save images to MinIO and get descriptions concurrently
            image_descriptions = await self._describe_images_concurrent_bytes(all_images, doc_id)

            # Send image descriptions and collect triples + entity contexts
            for idx, (img_info, description) in enumerate(zip(all_images, image_descriptions)):
                if description and description.strip():
                    image_text = f"{img_info['context']}\nDescription:\n{description}\n"
                    logger.info(f"Sending image {idx+1}/{len(all_images)} description to RAG")
                    r = TextDocument(
                        metadata=v.metadata,
                        text=image_text.encode("utf-8"),
                    )
                    await flow("output").send(r)
                    
                    # ---> Create IMAGE_CONTEXT and IMAGE_SOURCE triples for Graph RAG
                    image_uri = f"{doc_uri}/image_{idx + 1}"
                    image_uri_value = Value(value=image_uri, is_uri=True)
                    
                    # Add IMAGE_CONTEXT triple
                    image_triples.append(Triple(
                        s=image_uri_value,
                        p=Value(value=IMAGE_CONTEXT, is_uri=True),
                        o=Value(value=description, is_uri=False),
                    ))
                    
                    # Add IMAGE_SOURCE triple if we have MinIO path
                    # The path is constructed from doc_id and image index
                    image_path = f"minio://ocr-images/{self._safe_id(doc_id)}/image_{idx}.png"
                    image_triples.append(Triple(
                        s=image_uri_value,
                        p=Value(value=IMAGE_SOURCE, is_uri=True),
                        o=Value(value=image_path, is_uri=False),
                    ))
                    
                    # ---> Create EntityContext for graph embeddings (enables graph-retrieval to find this entity)
                    entity_contexts.append(EntityContext(
                        entity=image_uri_value,
                        context=description,
                    ))

        # ---> on_message > [triples flow] > emit IMAGE_CONTEXT/IMAGE_SOURCE triples for Graph RAG
        if image_triples:
            triples_msg = Triples(
                metadata=Metadata(
                    id=doc_id,
                    metadata=[],
                    user=v.metadata.user if hasattr(v.metadata, 'user') and v.metadata.user else "trustgraph",
                    collection=v.metadata.collection if hasattr(v.metadata, 'collection') and v.metadata.collection else "default",
                ),
                triples=image_triples,
            )
            await flow("triples").send(triples_msg)
            logger.info(f"Emitted {len(image_triples)} image context triples for document {doc_id}")

        # ---> on_message > [entity-contexts flow] > emit EntityContexts for graph embeddings (enables graph-retrieval)
        if entity_contexts:
            entity_contexts_msg = EntityContexts(
                metadata=Metadata(
                    id=doc_id,
                    metadata=[],
                    user=v.metadata.user if hasattr(v.metadata, 'user') and v.metadata.user else "trustgraph",
                    collection=v.metadata.collection if hasattr(v.metadata, 'collection') and v.metadata.collection else "default",
                ),
                entities=entity_contexts,
            )
            await flow("entity-contexts").send(entity_contexts_msg)
            logger.info(f"Emitted {len(entity_contexts)} entity contexts for graph embeddings for document {doc_id}")

        logger.info(f"DOCX extraction complete: {len(all_images)} images processed")

    # // ---> on_message > [_try_docx_extraction] > extract using python-docx with fallback
    def _try_docx_extraction(
        self, blob: bytes, doc_id: str, processed_hashes: Set[str]
    ) -> Tuple[bool, Optional[Dict], List[Dict]]:
        """
        Try to extract content using python-docx.
        Falls back to zipfile if python-docx fails completely.
        Images are stored in memory with bytes data for MinIO upload.
        """
        all_images = []
        structured_data = None

        try:
            # Load from bytes (in-memory)
            logger.debug("Trying python-docx load from bytes")
            doc = DocxDocument(BytesIO(blob))
            logger.info(f"python-docx loaded successfully: {len(doc.paragraphs)} paragraphs")

            # Extract all content (images stored as bytes)
            structured_data = self._extract_structured_docx(doc, doc_id, processed_hashes)
            logger.info(f"Extracted {len(structured_data['document']['paragraphs'])} paragraphs, "
                       f"{len(structured_data['document']['tables'])} tables")

            # Collect images with bytes data
            for img_info in structured_data["document"].get("images", []):
                all_images.append({
                    "data": img_info["data"],
                    "content_type": img_info["content_type"],
                    "context": f"Document Image - {img_info['filename']}",
                    "paragraph_index": img_info.get("paragraph_index")
                })

            return True, structured_data, all_images

        except Exception as e:
            logger.warning(f"python-docx load failed: {e}")
            return False, None, []

    # // ---> on_message > [_process_with_fallback] > handle fallback when extraction fails
    async def _process_with_fallback(self, blob: bytes, doc_id: str, v, flow):
        """Fallback processing when python-docx fails."""
        processed_hashes: Set[str] = set()
        images = self._extract_images_from_zip(blob, doc_id, processed_hashes)

        if images:
            logger.info(f"Fallback: Extracted {len(images)} images")
            all_images = [
                {
                    "data": img["data"],
                    "content_type": img["content_type"],
                    "context": f"Image {i+1} ({img['original_name']})",
                    "paragraph_index": None
                }
                for i, img in enumerate(images)
            ]
            descriptions = await self._describe_images_concurrent_bytes(all_images, doc_id)

            for img_info, desc in zip(all_images, descriptions):
                if desc:
                    r = TextDocument(
                        metadata=v.metadata,
                        text=f"{img_info['context']}\nDescription:\n{desc}\n".encode("utf-8")
                    )
                    await flow("output").send(r)
        else:
            logger.warning("Fallback extraction found no images")

    @staticmethod
    def add_args(parser):
        parser.add_argument(
            '--vllm-api-url',
            default='http://vllm-vision-server:8000/v1/chat/completions',
            help='vLLM API URL for image descriptions (default: http://vllm:8000/v1/chat/completions)'
        )
        parser.add_argument(
            '--vllm-model',
            default='google/gemma-3-4b-it',
            help='vLLM model name (default: Qwen/Qwen3-VL-4B-Instruct)'
        )
        FlowProcessor.add_args(parser)

    # // ---> [_safe_id] > sanitize directory name
    def _safe_id(self, value: str) -> str:
        if value is None:
            return "unknown"
        return re.sub(r"[^A-Za-z0-9._-]+", "_", str(value))

    # // ---> [_is_docx] > check if blob is a DOCX file
    def _is_docx(self, blob: bytes, content_type: Optional[str] = None) -> bool:
        """Check if blob is a DOCX file by verifying actual file content (magic bytes + ZIP structure)."""
        # DOCX files are ZIP archives starting with PK - ALWAYS check magic bytes first
        if len(blob) < 4 or blob[:4] != b'PK\x03\x04':
            # Not a ZIP file, so not a DOCX
            if content_type:
                logger.debug(f"File has content_type={content_type} but is not a ZIP archive (not DOCX)")
            return False

        # Further verify it's a DOCX by checking for word/ directory
        try:
            with zipfile.ZipFile(BytesIO(blob), 'r') as zf:
                names = zf.namelist()
                if any(name.startswith('word/') for name in names):
                    return True
        except Exception:
            pass

        return False

    # // ---> [_compute_image_hash] > MD5 hash for deduplication
    def _compute_image_hash(self, image_data: bytes) -> str:
        return hashlib.md5(image_data).hexdigest()

    # // ---> [_extract_images_from_zip] > fallback extraction via zipfile
    def _extract_images_from_zip(
        self, blob: bytes, doc_id: str, processed_hashes: Set[str]
    ) -> List[Dict[str, Any]]:
        """
        Extract images directly from DOCX ZIP archive.
        Returns list of dicts with image data, content_type, and original_name.
        Works even when python-docx can't parse the file.
        """
        extracted = []
        image_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp'}
        # EMF/WMF are Windows metafiles - we'll try to handle them
        metafile_extensions = {'.emf', '.wmf'}
        
        ext_to_content_type = {
            '.png': 'image/png', '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
            '.gif': 'image/gif', '.bmp': 'image/bmp', '.tiff': 'image/tiff',
            '.webp': 'image/webp'
        }

        try:
            with zipfile.ZipFile(BytesIO(blob), 'r') as zf:
                for name in zf.namelist():
                    # DOCX stores media in word/media/
                    if 'media/' not in name.lower():
                        continue

                    ext = Path(name).suffix.lower()
                    original_name = Path(name).name

                    if ext in image_extensions or ext in metafile_extensions:
                        try:
                            image_data = zf.read(name)

                            # Skip very large images
                            if len(image_data) > MAX_IMAGE_SIZE_MB * 1024 * 1024:
                                logger.debug(f"Skipping large image: {name}")
                                continue

                            img_hash = self._compute_image_hash(image_data)
                            if img_hash in processed_hashes:
                                continue
                            processed_hashes.add(img_hash)

                            # Try to convert EMF/WMF if PIL is available
                            if ext in metafile_extensions and PIL_AVAILABLE:
                                converted = self._try_convert_metafile(image_data)
                                if converted:
                                    image_data = converted
                                    ext = '.png'
                                else:
                                    # Skip unconvertible metafiles
                                    logger.debug(f"Skipping unconvertible metafile: {name}")
                                    continue

                            content_type = ext_to_content_type.get(ext, 'image/png')
                            extracted.append({
                                "data": image_data,
                                "content_type": content_type,
                                "original_name": original_name,
                                "filename": f"zip_{len(extracted)}{ext}"
                            })
                            logger.debug(f"Extracted from ZIP: {original_name}")

                        except Exception as e:
                            logger.debug(f"Failed to extract {name}: {e}")

        except zipfile.BadZipFile:
            logger.warning("Not a valid ZIP archive")
        except Exception as e:
            logger.warning(f"ZIP extraction failed: {e}")

        return extracted

    # // ---> [_try_convert_metafile] > convert EMF/WMF to PNG
    def _try_convert_metafile(self, data: bytes) -> Optional[bytes]:
        """Try to convert EMF/WMF metafile to PNG using PIL."""
        if not PIL_AVAILABLE:
            return None
        try:
            img = Image.open(BytesIO(data))
            output = BytesIO()
            img.save(output, format='PNG')
            return output.getvalue()
        except Exception:
            return None

    # // ---> [_extract_structured_docx] > parse DOCX with python-docx
    def _extract_structured_docx(
        self, doc, doc_id: str, processed_hashes: Set[str]
    ) -> Dict[str, Any]:
        """Extract all content from DOCX using python-docx. Images stored as bytes."""
        structured_data = {
            "document": {
                "paragraphs": [],
                "tables": [],
                "images": [],
                "properties": {}
            }
        }

        # Extract core properties
        try:
            props = doc.core_properties
            structured_data["document"]["properties"] = {
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "keywords": props.keywords or "",
                "comments": props.comments or "",
            }
        except Exception as e:
            logger.debug(f"Failed to extract core properties: {e}")

        image_counter = 0

        # Process paragraphs
        for para_idx, paragraph in enumerate(doc.paragraphs):
            try:
                para_data = self._extract_paragraph_data(
                    paragraph, para_idx, processed_hashes, image_counter
                )
                if para_data["text"] or para_data.get("images"):
                    structured_data["document"]["paragraphs"].append(para_data)

                # Collect images from paragraph
                for img_info in para_data.get("images", []):
                    img_info["paragraph_index"] = para_idx
                    structured_data["document"]["images"].append(img_info)
                    image_counter += 1

            except Exception as e:
                logger.debug(f"Failed to process paragraph {para_idx}: {e}")

        # Process tables
        for table_idx, table in enumerate(doc.tables):
            try:
                table_data = self._extract_table_data(table, table_idx)
                if table_data:
                    structured_data["document"]["tables"].append(table_data)

                # Extract images from table cells
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            try:
                                para_images = self._extract_paragraph_images(
                                    para, processed_hashes, image_counter
                                )
                                for img_info in para_images:
                                    img_info["table_index"] = table_idx
                                    structured_data["document"]["images"].append(img_info)
                                    image_counter += 1
                            except Exception:
                                pass

            except Exception as e:
                logger.debug(f"Failed to process table {table_idx}: {e}")

        # Also extract images from document's inline shapes (relationship-based extraction)
        try:
            doc_images = self._extract_document_images(doc, processed_hashes, image_counter)
            for img_info in doc_images:
                structured_data["document"]["images"].append(img_info)
        except Exception as e:
            logger.debug(f"Failed to extract document images: {e}")

        return structured_data

    # // ---> [_extract_paragraph_data] > extract content from a single paragraph
    def _extract_paragraph_data(
        self, paragraph, para_idx: int,
        processed_hashes: Set[str], image_counter: int
    ) -> Dict[str, Any]:
        """Extract text and images from a paragraph. Images stored as bytes."""
        para_data = {
            "paragraph_index": para_idx,
            "text": "",
            "style": None,
            "images": []
        }

        # Get paragraph style
        try:
            if paragraph.style:
                para_data["style"] = paragraph.style.name
        except Exception:
            pass

        # Get text content
        try:
            text = paragraph.text.strip()
            if text:
                para_data["text"] = text
        except Exception:
            pass

        # Extract inline images from runs (returns bytes)
        para_data["images"] = self._extract_paragraph_images(
            paragraph, processed_hashes, image_counter
        )

        return para_data

    # // ---> [_extract_paragraph_images] > extract images from paragraph runs
    def _extract_paragraph_images(
        self, paragraph, processed_hashes: Set[str], counter: int
    ) -> List[Dict]:
        """Extract images from paragraph's inline shapes. Returns bytes data."""
        images = []

        try:
            # Access inline shapes through the paragraph's XML
            for run in paragraph.runs:
                try:
                    # Check for inline shapes in the run's XML
                    inline_shapes = run._element.findall('.//' + '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
                    for inline in inline_shapes:
                        try:
                            # Get the blip element which contains the image reference
                            blip = inline.find('.//' + '{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                            if blip is not None:
                                # Get the relationship ID
                                embed_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                                rId = blip.get(embed_attr)
                                if rId:
                                    # Get the image part from the document's relationships
                                    try:
                                        part = paragraph.part.related_parts.get(rId)
                                        if part and hasattr(part, 'blob'):
                                            image_bytes = part.blob

                                            # Check size
                                            if len(image_bytes) > MAX_IMAGE_SIZE_MB * 1024 * 1024:
                                                continue

                                            # Check for duplicates
                                            img_hash = self._compute_image_hash(image_bytes)
                                            if img_hash in processed_hashes:
                                                continue
                                            processed_hashes.add(img_hash)

                                            # Determine content type and extension
                                            content_type = getattr(part, 'content_type', 'image/png')
                                            ext = self._content_type_to_ext(content_type)

                                            # Handle metafiles
                                            if ext in ['.emf', '.wmf'] and PIL_AVAILABLE:
                                                converted = self._try_convert_metafile(image_bytes)
                                                if converted:
                                                    image_bytes = converted
                                                    content_type = 'image/png'
                                                    ext = '.png'
                                                else:
                                                    continue

                                            filename = f"para_{counter + len(images)}{ext}"

                                            images.append({
                                                "filename": filename,
                                                "data": image_bytes,
                                                "content_type": content_type
                                            })
                                            logger.debug(f"Extracted inline image: {filename}")
                                    except Exception as e:
                                        logger.debug(f"Failed to get image part: {e}")
                        except Exception:
                            pass
                except Exception:
                    pass
        except Exception as e:
            logger.debug(f"Failed to extract paragraph images: {e}")

        return images

    # // ---> [_extract_document_images] > extract all images from document relationships
    def _extract_document_images(
        self, doc, processed_hashes: Set[str], counter: int
    ) -> List[Dict]:
        """Extract images using document part relationships as fallback. Returns bytes data."""
        images = []

        try:
            # Access all image parts from document relationships
            for rel in doc.part.rels.values():
                try:
                    if "image" in rel.reltype:
                        part = rel.target_part
                        if hasattr(part, 'blob'):
                            image_bytes = part.blob

                            # Check size
                            if len(image_bytes) > MAX_IMAGE_SIZE_MB * 1024 * 1024:
                                continue

                            # Check for duplicates
                            img_hash = self._compute_image_hash(image_bytes)
                            if img_hash in processed_hashes:
                                continue
                            processed_hashes.add(img_hash)

                            # Determine content type and extension
                            content_type = getattr(part, 'content_type', 'image/png')
                            ext = self._content_type_to_ext(content_type)

                            # Handle metafiles
                            if ext in ['.emf', '.wmf'] and PIL_AVAILABLE:
                                converted = self._try_convert_metafile(image_bytes)
                                if converted:
                                    image_bytes = converted
                                    content_type = 'image/png'
                                    ext = '.png'
                                else:
                                    continue

                            filename = f"doc_{counter + len(images)}{ext}"

                            images.append({
                                "filename": filename,
                                "data": image_bytes,
                                "content_type": content_type
                            })
                            logger.debug(f"Extracted document image: {filename}")

                except Exception as e:
                    logger.debug(f"Failed to extract image from relationship: {e}")

        except Exception as e:
            logger.debug(f"Failed to iterate document relationships: {e}")

        return images

    # // ---> [_content_type_to_ext] > map content type to file extension
    def _content_type_to_ext(self, content_type: str) -> str:
        """Map content type to file extension."""
        ext_map = {
            'image/png': '.png',
            'image/jpeg': '.jpg',
            'image/gif': '.gif',
            'image/bmp': '.bmp',
            'image/tiff': '.tiff',
            'image/webp': '.webp',
            'image/x-emf': '.emf',
            'image/x-wmf': '.wmf',
        }
        return ext_map.get(content_type, '.png')

    # // ---> [_extract_table_data] > extract table cell data
    def _extract_table_data(self, table, table_idx: int) -> Optional[Dict]:
        """Extract table data as structured dict."""
        try:
            rows_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    try:
                        # Get text from all paragraphs in cell
                        cell_texts = []
                        for para in cell.paragraphs:
                            para_text = para.text.strip()
                            if para_text:
                                cell_texts.append(para_text)
                        cell_text = " ".join(cell_texts)
                        row_data.append(cell_text)
                    except Exception:
                        row_data.append("")
                rows_data.append(row_data)

            if rows_data:
                return {
                    "table_index": table_idx,
                    "rows": len(rows_data),
                    "columns": len(rows_data[0]) if rows_data else 0,
                    "data": rows_data
                }
            return None

        except Exception as e:
            logger.debug(f"Table extraction failed: {e}")
            return None

    # // ---> [_format_document_header] > format document header
    def _format_document_header(self, structured_data: Dict, filename: str) -> str:
        """Format document header with metadata."""
        parts = []
        doc = structured_data.get("document", {})

        parts.append(f"# Word Document: {filename}")
        parts.append("")
        parts.append(f"Total Paragraphs: {len(doc.get('paragraphs', []))}")
        parts.append(f"Total Tables: {len(doc.get('tables', []))}")
        parts.append(f"Total Images: {len(doc.get('images', []))}")

        # Include core properties if available
        props = doc.get("properties", {})
        if props.get("title"):
            parts.append(f"Title: {props['title']}")
        if props.get("author"):
            parts.append(f"Author: {props['author']}")
        if props.get("subject"):
            parts.append(f"Subject: {props['subject']}")

        parts.append("")
        return "\n".join(parts)

    # // ---> [_format_document_text] > format document paragraphs as text
    def _format_document_text(self, structured_data: Dict) -> str:
        """Format document paragraphs as structured text for RAG."""
        parts = []
        doc = structured_data.get("document", {})

        parts.append("## Document Content:")
        parts.append("")

        current_style = None
        for para in doc.get("paragraphs", []):
            style = para.get("style")
            text = para.get("text", "")

            if not text:
                continue

            # Add style header if changed
            if style and style != current_style:
                if style.startswith("Heading"):
                    # Format headings appropriately
                    level = style.replace("Heading", "").strip()
                    try:
                        level_num = int(level)
                        parts.append("")
                        parts.append("#" * (level_num + 1) + f" {text}")
                        parts.append("")
                        current_style = style
                        continue
                    except ValueError:
                        pass
                current_style = style

            parts.append(text)
            parts.append("")

        return "\n".join(parts)

    # // ---> [_format_tables] > format tables as text
    def _format_tables(self, tables: List[Dict]) -> str:
        """Format tables as structured text."""
        parts = []

        parts.append("## Tables:")
        parts.append("")

        for table in tables:
            parts.append(f"Table {table['table_index'] + 1} ({table['rows']} rows × {table['columns']} columns):")
            parts.append("")
            for row in table.get("data", []):
                parts.append(" | ".join(str(c) if c else "" for c in row))
            parts.append("")

        return "\n".join(parts)

    # // ---> [_bytes_to_base64_data_url] > convert image bytes to base64 data URL
    def _bytes_to_base64_data_url(self, image_bytes: bytes, content_type: str = "image/png") -> str:
        """Convert image bytes to base64 data URL for vLLM."""
        b64 = base64.b64encode(image_bytes).decode("utf-8")
        return f"data:{content_type};base64,{b64}"

    # // ---> [_describe_images_concurrent_bytes] > concurrent vLLM requests with bytes
    async def _describe_images_concurrent_bytes(self, images: List[Dict], doc_id: str) -> List[str]:
        """
        Describe images using concurrent vLLM requests.
        Also saves images to MinIO storage.
        Processes up to MAX_CONCURRENT_VLLM_REQUESTS in parallel.
        """
        if not images:
            return []

        # Filter valid images (must have data)
        valid_images = [img for img in images if img.get("data")]
        if not valid_images:
            return [""] * len(images)

        logger.info(f"Processing {len(valid_images)} images concurrently")

        # Use get_running_loop() for Python 3.10+ compatibility
        try:
            loop = asyncio.get_running_loop()
        except RuntimeError:
            loop = asyncio.get_event_loop()

        with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_CONCURRENT_VLLM_REQUESTS) as executor:
            futures = [
                loop.run_in_executor(
                    executor,
                    self._describe_and_save_image,
                    img["data"],
                    img.get("content_type", "image/png"),
                    doc_id,
                    idx
                )
                for idx, img in enumerate(valid_images)
            ]
            results = await asyncio.gather(*futures, return_exceptions=True)

        # Process results
        descriptions = []
        for result in results:
            if isinstance(result, Exception):
                logger.error(f"Image description failed: {result}")
                descriptions.append("")
            else:
                descriptions.append(result or "")

        return descriptions

    # // ---> [_describe_and_save_image] > save to MinIO and describe via vLLM
    def _describe_and_save_image(
        self, image_bytes: bytes, content_type: str, doc_id: str, image_idx: int
    ) -> str:
        """Save image to MinIO and describe using vLLM."""
        try:
            # Determine extension from content type
            ext_map = {
                'image/png': '.png', 'image/jpeg': '.jpg', 'image/gif': '.gif',
                'image/bmp': '.bmp', 'image/webp': '.webp'
            }
            ext = ext_map.get(content_type, '.png')
            filename = f"image_{image_idx}{ext}"
            
            # Save to MinIO
            object_name = self.minio_storage.save_image(
                doc_id=doc_id,
                filename=filename,
                image_data=image_bytes,
                content_type=content_type
            )
            if object_name:
                logger.debug(f"Saved image to MinIO: {object_name}")
            
            # Describe via vLLM
            image_url = self._bytes_to_base64_data_url(image_bytes, content_type)

            payload = {
                "model": self.vllm_model,
                "messages": [{
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": (
                                "Describe this image from a document in detail. "
                                "Focus on: 1) Any text or labels visible, 2) Charts/graphs and their data, "
                                "3) Diagrams and their relationships, 4) Key visual elements. "
                                "Be concise but comprehensive."
                            )
                        },
                        {"type": "image_url", "image_url": {"url": image_url}}
                    ]
                }]
            }

            # Log the outgoing request (truncate base64 for readability)
            try:
                payload_log = json.loads(json.dumps(payload))
                try:
                    url_val = payload_log["messages"][0]["content"][1]["image_url"]["url"]
                    if url_val.startswith("data:"):
                        payload_log["messages"][0]["content"][1]["image_url"]["url"] = (
                            url_val[:60] + f"...[{len(url_val)} chars total]"
                        )
                except (KeyError, IndexError):
                    pass
                logger.debug(
                    "vLLM request: POST %s payload=%s",
                    self.vllm_api_url,
                    json.dumps(payload_log, ensure_ascii=False),
                )
            except Exception as log_err:
                logger.warning(f"Failed to serialize vLLM payload for logging: {log_err}")

            resp = requests.post(self.vllm_api_url, json=payload, timeout=VLLM_TIMEOUT_SECONDS)

            if resp.status_code != 200:
                logger.error(f"vLLM HTTP {resp.status_code}: {resp.text[:500]}")
                return f"Description unavailable (HTTP {resp.status_code})"

            data = resp.json()
            choice = (data.get("choices") or [{}])[0]
            content = (choice.get("message") or {}).get("content", "")

            if isinstance(content, str):
                logger.info(f"vLLM image description: {content.strip()}")
                return content.strip()
            if isinstance(content, list):
                return "\n".join(c.get("text", "") for c in content if c.get("type") == "text")
            return ""

        except requests.exceptions.Timeout:
            logger.error("vLLM request timed out")
            return "Description unavailable (timeout)"
        except requests.exceptions.ConnectionError as e:
            logger.error(f"vLLM connection error: {e}")
            return "Description unavailable (connection error)"
        except Exception as e:
            logger.error(f"vLLM error: {e}")
            return "Description unavailable"


def run():
    Processor.launch(default_ident, __doc__)
